#!/usr/bin/env python3
"""
Convertit un rapport circle_vs_faq_audit (Markdown) en DOCX avec diff visuel
natif Word (surlignage + barré) pour chaque article de la section 1
("Articles à mettre à jour").

Récupère `tchat_doc_audit_items` via Supabase REST pour calculer le diff
word-level entre `original_article_md` et `rewritten_article_md`, et rend le
tout via python-docx (w:highlight + w:strike natifs — pas de HTML inline).

Les sections 2/3/4 sont rendues en Markdown standard (pas de diff).

Usage:
  TCHAT_SUPPORT_SYNC_SUPABASE_URL=... TCHAT_SUPPORT_SYNC_SUPABASE_SERVICE_KEY=... \\
    python3 scripts/audit-to-docx.py <audit_run_id> <input.md> <output.docx>

L'input.md est le rapport brut exporté depuis `tchat_faq_documents`
(type=circle_vs_faq_audit). Pas besoin de le pré-patcher.
"""

import os
import sys
import re
import json
import difflib
import urllib.request
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt


def die(msg):
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(1)


SUPABASE_URL = os.environ.get("TCHAT_SUPPORT_SYNC_SUPABASE_URL")
SUPABASE_KEY = os.environ.get("TCHAT_SUPPORT_SYNC_SUPABASE_SERVICE_KEY")
if not SUPABASE_URL or not SUPABASE_KEY:
    die("Missing TCHAT_SUPPORT_SYNC_SUPABASE_URL or TCHAT_SUPPORT_SYNC_SUPABASE_SERVICE_KEY")

if len(sys.argv) != 4:
    die("Usage: python3 scripts/audit-to-docx.py <audit_run_id> <input.md> <output.docx>")

AUDIT_RUN_ID, INPUT_MD, OUTPUT_DOCX = sys.argv[1], sys.argv[2], sys.argv[3]


def sb_get(path):
    req = urllib.request.Request(
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"},
    )
    with urllib.request.urlopen(req) as r:
        return json.loads(r.read())


def fetch_update_items(run_id):
    out, offset = {}, 0
    while True:
        path = (
            f"tchat_doc_audit_items?audit_run_id=eq.{quote(run_id)}"
            f"&audit_type=eq.update_needed&select=article_url,article_name,analysis"
            f"&order=id&limit=1000&offset={offset}"
        )
        batch = sb_get(path)
        if not batch:
            break
        for row in batch:
            url = row.get("article_url")
            if not url:
                continue
            try:
                analysis = json.loads(row["analysis"])
            except (KeyError, json.JSONDecodeError):
                continue
            out[url] = {
                "original_md": analysis.get("original_article_md", ""),
                "rewritten_md": analysis.get("rewritten_article_md", ""),
            }
        if len(batch) < 1000:
            break
        offset += 1000
    return out


TYPO_NORMALIZE = str.maketrans({
    "\u2019": "'", "\u2018": "'",
    "\u201C": '"', "\u201D": '"',
    "\u2014": "-", "\u2013": "-",
    "\u2026": "...",
    "\u00A0": " ", "\u202F": " ",
})


def normalize(text):
    return text.translate(TYPO_NORMALIZE)


def tokenize(text):
    return re.findall(r"\S+|\s+", text)


def compute_diff_ops(old, new):
    """Returns list of (kind, text) tuples. kind ∈ {'equal','delete','insert'}."""
    old_toks, new_toks = tokenize(normalize(old)), tokenize(normalize(new))
    ops = []
    for op, i1, i2, j1, j2 in difflib.SequenceMatcher(a=old_toks, b=new_toks, autojunk=False).get_opcodes():
        if op == "equal":
            ops.append(("equal", "".join(new_toks[j1:j2])))
        elif op == "delete":
            ops.append(("delete", "".join(old_toks[i1:i2])))
        elif op == "insert":
            ops.append(("insert", "".join(new_toks[j1:j2])))
        elif op == "replace":
            ops.append(("delete", "".join(old_toks[i1:i2])))
            ops.append(("insert", "".join(new_toks[j1:j2])))
    return ops


# ---------- Markdown → DOCX (light parser) ----------

INLINE_RE = re.compile(
    r"(\*\*([^\*]+)\*\*|\*([^\*]+)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))"
)


def add_inline_runs(paragraph, text):
    """Parse bold/italic/code/link and emit runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2) is not None:
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif m.group(4) is not None:
            r = paragraph.add_run(m.group(4))
            r.font.name = "Consolas"
        elif m.group(5) is not None:
            r = paragraph.add_run(f"{m.group(5)} ({m.group(6)})")
            r.font.color.rgb = None
            r.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_diff_block(doc, original_md, rewritten_md):
    """Add a diff block with Word-native highlighting + strikethrough."""
    heading = doc.add_paragraph()
    heading.add_run("Diff visuel (proposée vs actuelle)").bold = True

    ops = compute_diff_ops(original_md, rewritten_md)

    paragraph = doc.add_paragraph()
    for kind, text in ops:
        for i, line in enumerate(text.split("\n")):
            if i > 0:
                paragraph = doc.add_paragraph()
            if not line:
                continue
            run = paragraph.add_run(line)
            if kind == "delete":
                run.font.highlight_color = WD_COLOR_INDEX.RED
                run.font.strike = True
            elif kind == "insert":
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN


URL_LINE_RE = re.compile(r"\*\*URL :\*\*\s*(\S+)")


def render_markdown(doc, md, diffs):
    """Render a Markdown string into a python-docx Document.

    Custom behavior: when a "#### Version proposée" block is encountered inside
    an article in section "Articles à mettre à jour", a colored diff block is
    injected right after (sourced from `diffs` by article URL).
    """
    lines = md.split("\n")
    i = 0
    in_update_section = False
    current_article_url = None
    diff_pending = False
    skip_details = False
    in_code_block = False

    def flush_diff():
        nonlocal diff_pending
        if diff_pending and current_article_url and current_article_url in diffs:
            d = diffs[current_article_url]
            add_diff_block(doc, d["original_md"], d["rewritten_md"])
        diff_pending = False

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = "Consolas"
            i += 1
            continue

        if line.strip().startswith("<details>"):
            flush_diff()
            skip_details = True
            i += 1
            continue
        if skip_details:
            if line.strip().startswith("</details>"):
                skip_details = False
            i += 1
            continue

        if line.strip().startswith("<") and ">" in line and not line.strip().startswith("<details>"):
            i += 1
            continue

        if line.startswith("## Articles à mettre à jour"):
            flush_diff()
            in_update_section = True
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        TOP_LEVEL_H2 = ("## Nouveaux articles proposés", "## Articles déjà alignés", "## Thèmes FAQ non couverts", "## Table des matières")
        if any(line.startswith(h) for h in TOP_LEVEL_H2):
            flush_diff()
            in_update_section = False
            current_article_url = None
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("# "):
            flush_diff()
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### 📝") or (line.startswith("### ") and not in_update_section):
            flush_diff()
            current_article_url = None
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=4)
            if in_update_section and heading_text == "Version proposée" and current_article_url and current_article_url in diffs:
                diff_pending = True
            i += 1
            continue

        url_match = URL_LINE_RE.search(line)
        if url_match and in_update_section:
            current_article_url = url_match.group(1)

        if line.startswith("---"):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            content = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, content)
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            add_inline_runs(p, line)
        else:
            doc.add_paragraph()
        i += 1


def main():
    print(f"Run ID: {AUDIT_RUN_ID}")
    print(f"Input:  {INPUT_MD}")
    print(f"Output: {OUTPUT_DOCX}")

    diffs = fetch_update_items(AUDIT_RUN_ID)
    print(f"Fetched {len(diffs)} update_needed items from Supabase")

    with open(INPUT_MD, "r", encoding="utf-8") as f:
        md = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    render_markdown(doc, md, diffs)
    doc.save(OUTPUT_DOCX)

    print(f"✅ Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
